from docx import Document
import os


def clean_filename(text):

    return (
        text.replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")
    )


def generate_resume_docx(
    content,
    company,
    role
):

    output_dir = "generated"

    os.makedirs(
        output_dir,
        exist_ok=True
    )

    filename = (
        f"{clean_filename(company)}_"
        f"{clean_filename(role)}.docx"
    )

    file_path = os.path.join(
        output_dir,
        filename
    )

    doc = Document()

    doc.add_heading(
        "Tailored Resume",
        level=1
    )

    doc.add_paragraph(
        content
    )

    doc.save(
        file_path
    )

    return file_path